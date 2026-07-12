"""문서 생성기 — AI가 도구로 호출하면 구조화 입력을 받아 docx/pdf/xlsx 바이트를 만든다.

- 순수 pip 의존성만 사용(system 라이브러리 불필요): python-docx / reportlab / openpyxl
- PDF 한글은 번들된 Pretendard TTF를 임베딩(뷰어에 폰트 없어도 안 깨짐).
  폰트 파일이 없으면 reportlab 내장 CID 폰트(HYSMyeongJo-Medium)로 폴백.
- 입력 마크다운은 최소 문법만 해석(# 제목, - / * 불릿, 1. 번호, 나머지는 문단)
"""
import io
import os
import re

_FONT_DIR = os.path.join(os.path.dirname(__file__), "assets", "fonts")

DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_CT = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
PDF_CT = "application/pdf"

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBER_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def _strip_inline(text: str) -> str:
    """굵게 등 인라인 마크다운 기호 제거(문서에는 평문으로)."""
    return _BOLD_RE.sub(r"\1", text)


def parse_blocks(markdown: str):
    """(kind, text, level) 블록 리스트로 변환. kind: h/p/bullet/number."""
    blocks = []
    for raw in (markdown or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        m = _HEADING_RE.match(line)
        if m:
            blocks.append(("h", _strip_inline(m.group(2)), len(m.group(1))))
            continue
        m = _BULLET_RE.match(line)
        if m:
            blocks.append(("bullet", _strip_inline(m.group(1)), 0))
            continue
        m = _NUMBER_RE.match(line)
        if m:
            blocks.append(("number", _strip_inline(m.group(1)), 0))
            continue
        blocks.append(("p", _strip_inline(line), 0))
    return blocks


# ---------------- DOCX ----------------

def _set_docx_korean_font(doc, name: str = "맑은 고딕"):
    """Normal 스타일에 한글(East Asian) 폰트를 명시해 뷰어별 대체 실패로 인한 깨짐을 방지."""
    from docx.oxml.ns import qn
    style = doc.styles["Normal"]
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def render_docx(title: str, markdown: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _set_docx_korean_font(doc)
    if title:
        doc.add_heading(title, level=0)
    for kind, text, level in parse_blocks(markdown):
        if kind == "h":
            doc.add_heading(text, level=min(max(level, 1), 3))
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------- PDF ----------------

_PDF_FONTS = {}


def _ensure_pdf_font():
    """한글 TTF(Pretendard)를 PDF에 임베딩해 어떤 뷰어에서도 깨지지 않게 한다.
    폰트 파일이 없으면 내장 CID 폰트로 폴백."""
    global _PDF_FONTS
    if _PDF_FONTS:
        return _PDF_FONTS
    from reportlab.pdfbase import pdfmetrics
    reg = os.path.join(_FONT_DIR, "Pretendard-Regular.ttf")
    bold = os.path.join(_FONT_DIR, "Pretendard-Bold.ttf")
    try:
        if not os.path.isfile(reg):
            raise FileNotFoundError(reg)
        from reportlab.pdfbase.ttfonts import TTFont
        pdfmetrics.registerFont(TTFont("KR", reg))
        pdfmetrics.registerFont(TTFont("KR-Bold", bold if os.path.isfile(bold) else reg))
        pdfmetrics.registerFontFamily("KR", normal="KR", bold="KR-Bold",
                                      italic="KR", boldItalic="KR-Bold")
        _PDF_FONTS = {"regular": "KR", "bold": "KR-Bold"}
    except Exception:
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
        _PDF_FONTS = {"regular": "HYSMyeongJo-Medium", "bold": "HYSMyeongJo-Medium"}
    return _PDF_FONTS


def render_pdf(title: str, markdown: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    from xml.sax.saxutils import escape

    fonts = _ensure_pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body_kr", parent=styles["Normal"], fontName=fonts["regular"], fontSize=11, leading=17)
    h = [ParagraphStyle(f"h{i}_kr", parent=styles["Heading%d" % i], fontName=fonts["bold"],
                        spaceBefore=10, spaceAfter=4) for i in (1, 2, 3)]
    title_style = ParagraphStyle("title_kr", parent=styles["Title"], fontName=fonts["bold"])

    buf = io.BytesIO()
    docp = SimpleDocTemplate(buf, pagesize=A4, topMargin=20 * mm, bottomMargin=20 * mm,
                             leftMargin=18 * mm, rightMargin=18 * mm)
    flow = []
    if title:
        flow.append(Paragraph(escape(title), title_style))
        flow.append(Spacer(1, 6))

    pending_items, pending_kind = [], None

    def flush_list():
        nonlocal pending_items, pending_kind
        if pending_items:
            flow.append(ListFlowable(
                [ListItem(Paragraph(escape(t), body)) for t in pending_items],
                bulletType="bullet" if pending_kind == "bullet" else "1"))
            pending_items, pending_kind = [], None

    for kind, text, level in parse_blocks(markdown):
        if kind in ("bullet", "number"):
            if pending_kind and pending_kind != kind:
                flush_list()
            pending_kind = kind
            pending_items.append(text)
            continue
        flush_list()
        if kind == "h":
            flow.append(Paragraph(escape(text), h[min(max(level, 1), 3) - 1]))
        else:
            flow.append(Paragraph(escape(text), body))
    flush_list()

    docp.build(flow)
    return buf.getvalue()


# ---------------- 회의록 (HP 표준 양식 오버레이) ----------------

_FORM_DIR = os.path.join(os.path.dirname(__file__), "assets", "forms")
_MINUTES_TEMPLATE = os.path.join(_FORM_DIR, "HP-QP-750-03_meeting.pdf")
_PAGE_H = 842.0  # A4 pt


def _mt_wrap(text: str, font: str, size: float, maxw: float) -> list:
    """폭 기준 줄바꿈(한글은 글자 단위). 명시적 개행도 처리."""
    from reportlab.pdfbase.pdfmetrics import stringWidth
    out, line = [], ""
    for ch in str(text or ""):
        if ch == "\n":
            out.append(line); line = ""; continue
        if stringWidth(line + ch, font, size) > maxw and line:
            out.append(line); line = ch
        else:
            line += ch
    if line:
        out.append(line)
    return out


def _as_str_list(v) -> list:
    """모델이 배열 대신 dict/숫자/문자열로 넘겨도 안전하게 문자열 리스트로 정규화."""
    if v is None:
        return []
    if isinstance(v, list):
        return [str(x) for x in v if x is not None and str(x).strip()]
    if isinstance(v, dict):
        return [str(x) for x in v.values() if x is not None and str(x).strip()]
    s = str(v).strip()
    return [s] if s else []


def _as_names(v) -> list:
    """참석자 이름 리스트로 정규화(문자열이면 쉼표/공백 분리)."""
    if isinstance(v, str):
        return [s.strip() for s in v.replace(",", " ").split() if s.strip()]
    return _as_str_list(v)


# 정적 양식은 런타임에 바뀌지 않으므로 바이트를 한 번만 읽어 캐시(반복 생성 시 디스크 I/O 절감).
_MINUTES_TEMPLATE_BYTES = None


def _minutes_template_bytes() -> bytes:
    global _MINUTES_TEMPLATE_BYTES
    if _MINUTES_TEMPLATE_BYTES is None:
        with open(_MINUTES_TEMPLATE, "rb") as f:
            _MINUTES_TEMPLATE_BYTES = f.read()
    return _MINUTES_TEMPLATE_BYTES


def render_minutes_pdf(fields: dict) -> bytes:
    """HP 표준 회의록 양식(HP-QP-750-03) 위에 값을 오버레이해 PDF 생성.
    양식은 정적 PDF(폼 필드 없음)이므로 좌표 기반으로 텍스트를 그려 원본 레이아웃을 그대로 유지한다."""
    import io as _io
    from reportlab.pdfgen import canvas
    from pypdf import PdfReader, PdfWriter

    fonts = _ensure_pdf_font()
    reg = fonts["regular"]
    buf = _io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(595, 842))
    c.setFillColorRGB(0.09, 0.09, 0.11)

    def line(text, x, ybase, size=9, center=None):
        if text is None or text == "":
            return
        c.setFont(reg, size)
        if center is not None:
            c.drawCentredString(center, _PAGE_H - ybase, str(text))
        else:
            c.drawString(x, _PAGE_H - ybase, str(text))

    def block(text, x, ytop, maxw, maxlines, size=9.5, lh=14):
        if text is None or text == "":
            return
        c.setFont(reg, size)
        lines = _mt_wrap(text, reg, size, maxw)
        if len(lines) > maxlines:  # 1페이지 양식 한계 초과 시 잘림을 명시(무단 누락 방지)
            lines = lines[:maxlines]
            lines[-1] = lines[-1][:max(0, len(lines[-1]) - 6)] + " …(이하 생략)"
        y = ytop
        for ln in lines:
            c.drawString(x, _PAGE_H - y, ln); y += lh

    g = fields.get
    line(g("mgmt_no"), 75, 133, 8.5)
    line(g("approval_no"), 360, 133, 8.5)
    line(g("datetime"), 125, 155)
    line(g("dept"), 326, 155)
    line(g("writer"), 486, 155)
    line(g("place"), 125, 178)
    line(g("ext_company"), 326, 178)
    inside = _as_names(g("inside"))
    outside = _as_names(g("outside"))
    for i, cx in enumerate([143, 190, 238]):
        if i < len(inside):
            line(inside[i], 0, 201, 8, center=cx)
    for i, cx in enumerate([353, 409, 457, 504, 552]):
        if i < len(outside):
            line(outside[i], 0, 201, 8, center=cx)
    line(g("purpose"), 125, 225)
    block(g("content"), 125, 246, 446, 18, size=9.5, lh=14)
    tops = [513, 542, 572, 601, 631, 660, 690, 719, 749]
    notes = _as_str_list(g("notes"))
    for i, item in enumerate(notes[:9]):
        block(item, 124, tops[i], 447, 2, size=9, lh=13.5)
    c.save(); buf.seek(0)

    overlay = PdfReader(buf).pages[0]
    tmpl = PdfReader(_io.BytesIO(_minutes_template_bytes()))
    page = tmpl.pages[0]
    page.merge_page(overlay)
    writer = PdfWriter(); writer.add_page(page)
    out = _io.BytesIO(); writer.write(out)
    return out.getvalue()


# ---------------- 회사 표준 양식(AcroForm 채우기) ----------------
# 견적서/거래명세서 등은 필드 이름이 박힌 AcroForm PDF다. 좌표 오버레이 대신
# PyMuPDF로 필드 값을 채우면(내장 CID 한글 폰트로 외관 스트림까지 생성) 어떤 뷰어에서도
# 값이 그대로 보인다. 회사 원본 레이아웃/도장/QR 필드는 손대지 않고 유지된다.

_ESTIMATE_TEMPLATE = os.path.join(_FORM_DIR, "estimate.pdf")
_ESTIMATE_SIMPLE_TEMPLATE = os.path.join(_FORM_DIR, "estimate_simple.pdf")
_STATEMENT_TEMPLATE = os.path.join(_FORM_DIR, "transaction_statement.pdf")
_OVERTIME_TEMPLATE = os.path.join(_FORM_DIR, "overtime_request.pdf")
_EXPENSE_TEMPLATE = os.path.join(_FORM_DIR, "expense_report.pdf")
_LEAVE_TEMPLATE = os.path.join(_FORM_DIR, "leave_request.pdf")
_FLEX_TEMPLATE = os.path.join(_FORM_DIR, "flexible_work.pdf")

# 콤보박스 허용값(원본 양식에 내장된 선택지 그대로). 앞뒤 공백까지 원본과 일치해야 채워진다.
OT_DEPARTMENTS = [" 경영지원", "기술엔지니어링", "연구개발전담부서"]
EXPENSE_PROOFS = ["간이영수증", "기타", "분실", "세금계산서", "영수증", "카드영수증", "현금영수증"]
LEAVE_DEPARTMENTS = [" ", "연구개발전담부서", "경영지원", "기술엔지니어링"]
LEAVE_TYPES = [" ", "연차", "반차", "반반차", "경조사", "병가", "공가", "특별휴가", "무급휴가"]

_FORM_TEMPLATE_BYTES: dict = {}


def _form_template_bytes(path: str) -> bytes:
    """정적 양식 PDF 바이트를 1회만 읽어 캐시(반복 생성 시 디스크 I/O 절감)."""
    if path not in _FORM_TEMPLATE_BYTES:
        with open(path, "rb") as f:
            _FORM_TEMPLATE_BYTES[path] = f.read()
    return _FORM_TEMPLATE_BYTES[path]


def _won(n) -> str:
    """금액을 천단위 콤마 문자열로. 숫자 변환 불가하면 빈 문자열."""
    try:
        return f"{int(round(float(n))):,}"
    except (TypeError, ValueError):
        return ""


def _num(v):
    """콤마·통화기호 섞인 값도 float로. 불가하면 None."""
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = re.sub(r"[^0-9.\-]", "", str(v))
    try:
        return float(s) if s not in ("", "-", ".") else None
    except ValueError:
        return None


# 열 성격별 정렬: 금액/단가/합계=우측, 단위/수량=가운데, 그 외=좌측.
_ACRO_RIGHT = re.compile(r"단가|공[급금]가액|금\s?액|합계|\*10%")
_ACRO_CENTER = re.compile(r"단위|수량")
_FILL_RGB = (0.09, 0.09, 0.11)


def _acro_align(name: str) -> str:
    if _ACRO_RIGHT.search(name):
        return "r"
    if _ACRO_CENTER.search(name):
        return "c"
    return "l"


def _subset_pretendard(chars) -> bytes | None:
    """주어진 문자들만 담아 Pretendard를 서브셋한 폰트 바이트를 반환(없거나 실패 시 None).
    전체 폰트 임베딩은 수 MB라 그릴 글자만 골라 넣는다. PyMuPDF subset_fonts()는
    양식 원본 CID 폰트를 손상시켜 쓰지 않는다."""
    reg = os.path.join(_FONT_DIR, "Pretendard-Regular.ttf")
    if not os.path.isfile(reg):
        return None
    try:
        import io as _io
        from fontTools import subset as _ftsubset
        from fontTools.ttLib import TTFont as _TTFont
        tt = _TTFont(reg)
        ss = _ftsubset.Subsetter()
        ss.populate(unicodes=sorted({ord(c) for c in chars}))
        ss.subset(tt)
        _b = _io.BytesIO(); tt.save(_b)
        return _b.getvalue()
    except Exception:  # noqa: BLE001 — 서브셋 실패 시 폴백은 호출측에서
        return None


def _fill_acroform(template_path: str, values: dict, overlays=None) -> bytes:
    """AcroForm 필드를 이름 기준으로 채워 PDF 바이트 반환.

    폼 내장 폰트(AdobeGothicStd)는 ASCII/숫자/공백을 전각으로 렌더해 '띄어쓰기가 벌어져'
    보이고, PyMuPDF 위젯 API로는 정렬(Q)·커스텀 폰트 지정이 불안정하다. 그래서 각 필드의
    사각형만 읽어 값을 직접 Pretendard로 그리고(자연스러운 자간·열별 정렬), 입력 위젯은
    평탄화(삭제)해 어떤 뷰어에서도 동일하게 보이게 한다. 도장/QR/서명(버튼) 필드는 보존.
    - 텍스트/콤보: 값 그리기(금액=우측, 단위/수량=가운데, 그 외=좌측)
    - 체크박스: 값이 참이면 벡터 체크표시
    overlays: [{page, text, x|right, y, size}] — 폼 필드가 없는 자리(예: 합계)에 텍스트를 그린다."""
    import fitz  # PyMuPDF (지연 임포트)

    reg = os.path.join(_FONT_DIR, "Pretendard-Regular.ttf")
    has_font = os.path.isfile(reg)

    # 그릴 문자만 모아 Pretendard를 서브셋 → 임베딩.
    chars = set("0123456789.,-/():%원₩ ")
    for v in values.values():
        if isinstance(v, str):
            chars.update(v)
    for ov in (overlays or []):
        chars.update(str(ov.get("text") or ""))
    fontbuffer = _subset_pretendard(chars) if has_font else None

    if fontbuffer:
        measure = fitz.Font(fontbuffer=fontbuffer)
    elif has_font:
        measure = fitz.Font(fontfile=reg)
    else:
        measure = fitz.Font("helv")
    fontname = "pret" if has_font else "helv"
    _ready = set()   # 폰트 등록된 페이지 인덱스

    def _ensure_font(page, pi):
        if pi in _ready or not has_font:
            return
        if fontbuffer:
            page.insert_font(fontname="pret", fontbuffer=fontbuffer)
        else:
            page.insert_font(fontname="pret", fontfile=reg)
        _ready.add(pi)

    def _draw(page, pi, rect, text, size, align):
        text = str(text)
        if not text:
            return
        _ensure_font(page, pi)
        size = max(6.0, min(float(size or 9.0), 15.0))
        tw = measure.text_length(text, size)
        avail = rect.width - 6
        if tw > avail and tw > 0:                # 칸을 넘치면 폰트를 줄여 맞춤
            size = max(6.0, size * avail / tw)
            tw = measure.text_length(text, size)
        if align == "r":
            x = rect.x1 - tw - 3
        elif align == "c":
            x = rect.x0 + (rect.width - tw) / 2
        else:
            x = rect.x0 + 3
        y = rect.y0 + (rect.height + size * 0.70) / 2   # 세로 중앙
        page.insert_text((x, y), text, fontsize=size, fontname=fontname, color=_FILL_RGB)

    def _check(page, rect):
        w, h = rect.width, rect.height
        sh = page.new_shape()
        sh.draw_line(fitz.Point(rect.x0 + w * 0.20, rect.y0 + h * 0.52),
                     fitz.Point(rect.x0 + w * 0.42, rect.y0 + h * 0.74))
        sh.draw_line(fitz.Point(rect.x0 + w * 0.42, rect.y0 + h * 0.74),
                     fitz.Point(rect.x0 + w * 0.80, rect.y0 + h * 0.24))
        sh.finish(width=max(1.0, h * 0.10), color=_FILL_RGB)
        sh.commit()

    doc = fitz.open(stream=_form_template_bytes(template_path), filetype="pdf")
    try:
        # 1) 위젯 사각형/값을 먼저 수집(삭제 후에도 그리기 위해 rect 복사)
        plan = []            # (page_index, rect, kind, payload)
        input_types = {"Text", "ComboBox", "ListBox", "CheckBox"}
        for pi, page in enumerate(doc):
            for w in (page.widgets() or []):
                ftype = (w.field_type_string or "")
                if ftype not in input_types:
                    continue          # 도장/QR/서명 등은 건드리지 않음
                v = values.get(w.field_name)
                if ftype == "CheckBox":
                    if v:
                        plan.append((pi, fitz.Rect(w.rect), "check", None))
                elif v not in (None, ""):
                    plan.append((pi, fitz.Rect(w.rect), "text",
                                 (str(v), w.text_fontsize, _acro_align(w.field_name))))
        # 2) 입력 위젯 평탄화(빈 칸 포함 제거 → 인쇄/뷰어 일관)
        for page in doc:
            for w in (page.widgets() or []):
                if (w.field_type_string or "") in input_types:
                    try:
                        page.delete_widget(w)
                    except Exception:  # noqa: BLE001
                        pass
        # 3) 값 그리기
        for pi, rect, kind, payload in plan:
            page = doc[pi]
            if kind == "check":
                _check(page, rect)
            else:
                text, size, align = payload
                _draw(page, pi, rect, text, size, align)
        # 4) 폼 필드 없는 자리(합계 등)에 우측정렬 텍스트
        for ov in (overlays or []):
            page = doc[ov.get("page", 0)]
            text = str(ov.get("text") or "")
            if not text:
                continue
            size = ov.get("size", 10)
            _ensure_font(page, ov.get("page", 0))
            tw = measure.text_length(text, size)
            x = ov.get("x")
            if x is None and ov.get("right") is not None:
                x = ov["right"] - tw
            page.insert_text((x or 0, ov["y"]), text, fontsize=size, fontname=fontname,
                             color=_FILL_RGB)
        out = doc.tobytes(garbage=4, deflate=True)
    finally:
        doc.close()
    return out


def _combo_match(value, options):
    """콤보박스 값을 원본 선택지로 정규화(공백 무시 비교). 못 맞추면 None."""
    s = str(value or "").strip()
    if not s:
        return None
    for o in options:
        if o.strip() == s:
            return o
    return None


def _line_items(items, max_rows: int):
    """[{name,unit,qty,price,amount,note}] → 행별 정규화 + 공급가액/세액 자동계산.
    반환: (rows, 공급가액합계, 세액합계). amount 미지정 시 수량*단가로 계산."""
    rows, sup_total, tax_total = [], 0.0, 0.0
    for it in (items or [])[:max_rows]:
        if not isinstance(it, dict):
            it = {"name": str(it)}
        qty, price = _num(it.get("qty")), _num(it.get("price"))
        amount = _num(it.get("amount"))
        if amount is None and qty is not None and price is not None:
            amount = qty * price
        tax = round(amount * 0.1) if amount is not None else None
        if amount is not None:
            sup_total += amount
        if tax is not None:
            tax_total += tax
        rows.append({
            "name": it.get("name"), "unit": it.get("unit"),
            "qty": it.get("qty"), "price": price, "amount": amount, "tax": tax,
            "note": it.get("note"),
        })
    return rows, sup_total, tax_total


def render_estimate_pdf(data: dict) -> bytes:
    """HP 견적서(또는 간이견적서) AcroForm을 채워 PDF 생성.
    품목의 공급가액/세액/합계는 수량·단가로 자동계산한다."""
    g = data.get
    template = _ESTIMATE_SIMPLE_TEMPLATE if data.get("simple") else _ESTIMATE_TEMPLATE
    rows, sup, tax = _line_items(g("items"), 11)
    vals = {
        "거래처이름": g("client_name"),
        "거래처 담당자 성함": g("client_manager"),
        "거래처 담당자 연락처": g("client_contact"),
        "거래처 담당자혹은 거래처 팩스번호": g("client_fax"),
        "견적 이름": g("estimate_name"),
        "시험 대상품목 명 혹은 규격(규정)": g("target_spec"),
        "HP엔지니어링 담다자 성함": g("hp_manager"),
        "HP엔지니어링 담당자 사내직통번호": g("hp_phone"),
        "HP엔지니어링 담당자 이메일 주소": g("hp_email"),
        "견적번호 기입필드": g("estimate_no"),
        "간이견적서발행일": g("issue_date"),   # 견적서 발행일(양식상 필드명이 이러함)
        "임시 작업완료기한": g("work_deadline"),
        "인도장소 기입란": g("delivery_place"),
        "추가 협의된 사항이나, 확인할 사항, 명시할 사항 기입란": g("confirm_notes"),
        "대금 결제일": g("payment_date"),
        "공급가액의 합계액": _won(sup),
        "공급가액의 합계액*10%": _won(tax),
        "세액포함 총 합계액 기입필드": _won(sup + tax),
    }
    for i, r in enumerate(rows, 1):
        # 원본 양식 오타: 7행 공급가액 필드명이 '공금가액7'
        sup_key = "공금가액7" if i == 7 else f"공급가액{i}"
        vals.update({
            f"견적품명 {i}": r["name"], f"단위{i}": r["unit"], f"수량{i}": r["qty"],
            f"단가{i}": _won(r["price"]) if r["price"] is not None else None,
            sup_key: _won(r["amount"]) if r["amount"] is not None else None,
            f"공급가액*10%_{i}": _won(r["tax"]) if r["tax"] is not None else None,
            f"비고{i}": r["note"],
        })
    return _fill_acroform(template, vals)


def render_transaction_statement_pdf(data: dict) -> bytes:
    """HP 거래명세서 AcroForm을 채워 PDF 생성. 공급가액/세액/합계 자동계산."""
    g = data.get
    rows, sup, tax = _line_items(g("items"), 15)
    vals = {
        "거래처 이름": g("client_name"),
        "거래처 담당자 이름": g("client_manager"),
        "작업명(시험명)": g("work_name"),
        "거래명세번호(거래명세송부일-거래처코드-거래명세송부횟수(오름차순)": g("statement_no"),
        "작업완료 현황(시기)기입": g("work_done"),
        "인도완료 일자(현황)기입": g("delivery_done"),
        "거래명세서 첨부일 기입": g("attach_date"),
        "대금지불 조건 기입(ex, 익월 20일지급, 거래명세 교부 당일지급 등)": g("payment_terms"),
        "공급가액 합계액": _won(sup),
        "공급가액*10% 합계액": _won(tax),
        "합계금액(공급가액+공급가액*10%)": _won(sup + tax),
    }
    for i, r in enumerate(rows, 1):
        vals.update({
            f"거래일자(시험일자)혹은 품명 등{i}": r["name"], f"단위{i}": r["unit"],
            f"수량{i}": r["qty"],
            f"단가{i}": _won(r["price"]) if r["price"] is not None else None,
            f"공급가액{i}": _won(r["amount"]) if r["amount"] is not None else None,
            f"공급가액*10%_{i}": _won(r["tax"]) if r["tax"] is not None else None,
            f"비고{i}": r["note"],
        })
    return _fill_acroform(_STATEMENT_TEMPLATE, vals)


def _hhmm(t):
    """'18:30' / '18시30분' / '1830' → (시, 분). 파싱 실패하면 None."""
    m = re.match(r"\s*(\d{1,2})\s*[:시]?\s*(\d{2})?\s*분?\s*$", str(t or ""))
    if not m:
        return None
    h = int(m.group(1))
    mm = int(m.group(2) or 0)
    if h > 23 or mm > 59:
        return None
    return h, mm


def render_overtime_request_pdf(data: dict) -> bytes:
    """HP 연장 근로 신청서 AcroForm을 채워 PDF 생성.
    총 연장근로시간은 시작/종료 시간으로 자동계산(명시값이 있으면 그대로 사용)."""
    g = data.get
    total = g("total_hours")
    if not total:
        a, b = _hhmm(g("start_time")), _hhmm(g("end_time"))
        if a and b:
            mins = (b[0] * 60 + b[1]) - (a[0] * 60 + a[1])
            if mins > 0:
                total = f"{mins // 60}시간" + (f" {mins % 60}분" if mins % 60 else "")
    vals = {
        "신청일자": g("apply_date"), "신청자": g("applicant"), "입사일": g("hire_date"),
        "부서": _combo_match(g("dept"), OT_DEPARTMENTS), "직위": g("position"), "사번": g("emp_no"),
        "관리번호": g("mgmt_no"), "승인번호": g("approval_no"),
        "작성일자": g("write_date"), "결제일자": g("approve_date"),
        "연장 근무 시작시간": g("start_time"),
        "연장 근무 종료시간(최대 4시간이상 X)": g("end_time"),
        "총 연장 근로시간": total,
        "연장근무 사유": g("reason"), "특이 사항": g("remarks"),
        "작성자 성명": g("writer") or g("applicant"),
    }
    return _fill_acroform(_OVERTIME_TEMPLATE, vals)


def render_expense_report_pdf(data: dict) -> bytes:
    """HP 지출결의서 AcroForm을 채워 PDF 생성. 지출항목 최대 8행, 금액 합계는 자동계산해
    합계 칸(폼 필드 없음)에 오버레이한다. 지급방법 체크박스/증빙구분 콤보 처리."""
    g = data.get
    vals = {
        "작성자성명": g("writer_name"), "작성자직급": g("writer_rank"),
        "작성자사번": g("writer_empno"), "기안일자": g("draft_date"),
        "관리번호": g("mgmt_no"), "승인번호": g("approval_no"),
        "지출 제목": g("title"), "지출 사유": g("reason"),
        "급여수령계좌은행": g("bank"), "급여수령계좌번호": g("account"),
        "작성자 성명": g("writer_name"),  # 예금주 자리
        "특기사항_증빙구분이 기타, 분실일 경우 사유 필수 작성": g("special_notes"),
    }
    total = 0.0
    for i, it in enumerate((g("items") or [])[:8]):
        if not isinstance(it, dict):
            it = {"item": str(it)}
        j = i  # 항목명은 1부터(지출항목1), 나머지 열은 0부터(.0)로 이름이 다름
        amt = _num(it.get("amount"))
        if amt is not None:
            total += amt
        vals[f"지출항목{i + 1}"] = it.get("item")
        vals[f"(구입처,지출처)거래처.{j}"] = it.get("vendor")
        vals[f"적요,기타사항,계정과목.{j}"] = it.get("memo")
        vals[f"금 액.{j}"] = _won(amt) if amt is not None else None
        vals[f"증빙구분.{j}"] = _combo_match(it.get("proof"), EXPENSE_PROOFS)
        vals[f"관련프로젝트.{j}"] = it.get("project")
        vals[f"비고.{j}"] = it.get("note")
    # 지급방법 체크박스(하나 선택)
    cbmap = {"계좌이체": "계좌이체체크박스", "현금": "현금수령희망시 체크박스",
             "법인카드": "법인카드 지출시 체크박스"}
    pay_cb = cbmap.get(str(g("pay_method") or "").strip())
    if pay_cb:
        vals[pay_cb] = True
    overlays = []
    if total > 0:  # 합계 칸엔 폼 필드가 없어 숫자를 우측정렬로 그린다("원(부가세 포함)" 앞)
        overlays.append({"page": 0, "text": _won(total), "right": 424, "y": 521, "size": 10})
    return _fill_acroform(_EXPENSE_TEMPLATE, vals, overlays)


def render_leave_request_pdf(data: dict) -> bytes:
    """HP 휴가신청서 AcroForm을 채워 PDF 생성. 부서/휴가종류 콤보, 오전·오후 반차 체크박스.
    발생/사용/잔여 연차를 주면 소진율(사용/발생)은 비어 있을 때 자동계산한다."""
    g = data.get
    # 소진율 자동계산(명시값 없을 때). 발생 연차 대비 사용 연차 백분율.
    rate = g("usage_rate")
    if not rate:
        grant, used = _num(g("granted_days")), _num(g("used_days"))
        if grant and used is not None:
            rate = f"{round(used / grant * 100)}%"
    vals = {
        "신청일자": g("apply_date"), "신청자": g("applicant"), "입사일": g("hire_date"),
        "부서": _combo_match(g("dept"), LEAVE_DEPARTMENTS), "직위": g("position"), "사번": g("emp_no"),
        "휴가 종류": _combo_match(g("leave_type"), LEAVE_TYPES),
        "시간 기간": g("start"), "종료시간": g("end"), "휴가 일수": g("days"),
        "발생 연차": g("granted_days"), "사용 연차": g("used_days"),
        "잔여연차": g("remaining_days"), "소진율": rate,
        "휴가 사유": g("reason"), "연락처": g("contact"),
        "업무대리자": g("deputy"), "인수인계 사항": g("handover"),
        "이름": g("applicant"),   # 하단 작성자 이름
    }
    # 오전/오후 반차 체크박스(반차·반반차 선택 시 사용)
    half = str(g("half_day") or "")
    if "오전" in half:
        vals["Check Box6"] = True
    if "오후" in half:
        vals["Check Box7"] = True
    return _fill_acroform(_LEAVE_TEMPLATE, vals)


def render_flexible_work_pdf(data: dict) -> bytes:
    """HP 유연근무 신청서 생성. 이 양식은 입력 필드가 없는 '평면' PDF라 표 셀 좌표에
    Pretendard로 값을 직접 그린다(줄바꿈은 셀 안에서 자동 줄맞춤). 원본 레이아웃은 보존."""
    import fitz  # PyMuPDF (지연 임포트)

    g = data.get
    # (rect, text, size, align, 세로중앙여부) — 좌표는 원본 표 격자에서 추출(612x792, pt)
    def R(x0, y0, x1, y1):
        return fitz.Rect(x0, y0, x1, y1)
    cells = [
        (R(230, 126, 331, 150), g("name"), 10, "center", True),
        (R(444, 126, 591, 150), g("dept"), 10, "center", True),
        (R(230, 152, 331, 177), g("position"), 10, "center", True),
        (R(444, 152, 591, 177), g("emp_no"), 10, "center", True),
        (R(126, 179, 588, 204), g("apply_date"), 10, "left", True),
        (R(126, 208, 588, 360), g("reason"), 10, "left", False),
        (R(126, 369, 588, 735), g("work_plan"), 10, "left", False),
        (R(126, 741, 588, 754), g("remarks"), 8, "left", True),
        (R(112, 106, 360, 122), g("mgmt_no"), 9, "left", True),
        (R(404, 106, 588, 122), g("approval_no"), 9, "left", True),
    ]
    chars = set(" 0123456789.,-/():%")
    for _r, t, _s, _a, _c in cells:
        if t:
            chars.update(str(t))
    fontbuffer = _subset_pretendard(chars)
    reg = os.path.join(_FONT_DIR, "Pretendard-Regular.ttf")
    has_font = os.path.isfile(reg)
    align_map = {"left": 0, "center": 1, "right": 2}

    doc = fitz.open(stream=_form_template_bytes(_FLEX_TEMPLATE), filetype="pdf")
    try:
        page = doc[0]
        if has_font:
            if fontbuffer:
                page.insert_font(fontname="pret", fontbuffer=fontbuffer)
            else:
                page.insert_font(fontname="pret", fontfile=reg)
        fontname = "pret" if has_font else "helv"
        for rect, text, size, align, vcenter in cells:
            text = str(text or "")
            if not text:
                continue
            single = "\n" not in text and len(text) < 40
            if vcenter and single:
                # 한 줄 값은 셀 세로 중앙에 직접 그린다(작은 칸에서 상단 쏠림 방지)
                mfont = fitz.Font(fontbuffer=fontbuffer) if fontbuffer else (
                    fitz.Font(fontfile=reg) if has_font else fitz.Font("helv"))
                sz = size
                tw = mfont.text_length(text, sz)
                avail = rect.width - 6
                if tw > avail and tw > 0:
                    sz = max(6.0, sz * avail / tw); tw = mfont.text_length(text, sz)
                if align == "center":
                    x = rect.x0 + (rect.width - tw) / 2
                elif align == "right":
                    x = rect.x1 - tw - 3
                else:
                    x = rect.x0 + 3
                y = rect.y0 + (rect.height + sz * 0.70) / 2
                page.insert_text((x, y), text, fontsize=sz, fontname=fontname, color=_FILL_RGB)
            else:
                # 여러 줄/긴 값은 셀 안에서 자동 줄바꿈(넘치면 폰트 축소 재시도)
                sz = size
                while sz >= 6.0:
                    rc = page.insert_textbox(rect, text, fontsize=sz, fontname=fontname,
                                             color=_FILL_RGB, align=align_map.get(align, 0))
                    if rc >= 0:
                        break
                    sz -= 0.5
        out = doc.tobytes(garbage=4, deflate=True)
    finally:
        doc.close()
    return out


# ---------------- XLSX ----------------

def render_xlsx(sheets: list) -> bytes:
    """sheets: [{name, headers:[str], rows:[[cell,...]]}]"""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    wb.remove(wb.active)
    if not sheets:
        sheets = [{"name": "Sheet1", "headers": [], "rows": []}]
    for i, sheet in enumerate(sheets):
        name = str(sheet.get("name") or f"Sheet{i+1}")[:31] or f"Sheet{i+1}"
        ws = wb.create_sheet(title=name)
        headers = sheet.get("headers") or []
        if headers:
            ws.append([str(x) for x in headers])
            for c in ws[1]:
                c.font = Font(bold=True)
        for row in (sheet.get("rows") or []):
            ws.append(["" if v is None else v for v in row])
        # 열 너비 자동(대략)
        for col in ws.columns:
            width = max((len(str(c.value)) if c.value is not None else 0) for c in col)
            ws.column_dimensions[col[0].column_letter].width = min(max(width + 2, 8), 60)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def render(fmt: str, *, title: str = "", markdown: str = "", sheets: list | None = None):
    """포맷별 (bytes, content_type, 확장자) 반환."""
    fmt = (fmt or "").lower()
    if fmt == "docx":
        return render_docx(title, markdown), DOCX_CT, "docx"
    if fmt == "pdf":
        return render_pdf(title, markdown), PDF_CT, "pdf"
    if fmt == "xlsx":
        return render_xlsx(sheets or []), XLSX_CT, "xlsx"
    raise ValueError(f"지원하지 않는 형식: {fmt}")
